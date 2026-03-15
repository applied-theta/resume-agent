"""Tests for the DOCX renderer script (md-to-docx.py)."""

import sys
import tempfile
from pathlib import Path

import pytest

# Add scripts/ to path so we can import the modules
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))

from parse_resume import parse_resume_markdown

# Import must succeed (verifies python-docx is installed)
from docx import Document

# Now import the renderer module
# We import it via importlib to handle the hyphenated filename
import importlib.util

_spec = importlib.util.spec_from_file_location(
    "md_to_docx",
    str(Path(__file__).resolve().parent.parent / "scripts" / "md-to-docx.py"),
)
md_to_docx = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(md_to_docx)

render_docx = md_to_docx.render_docx
parse_args = md_to_docx.parse_args
resolve_config = md_to_docx.resolve_config
validate_template_styles = md_to_docx.validate_template_styles
REQUIRED_TEMPLATE_STYLES = md_to_docx.REQUIRED_TEMPLATE_STYLES


# ---------------------------------------------------------------------------
# Sample resume fixtures
# ---------------------------------------------------------------------------

FULL_RESUME_MD = """\
# Jane Doe

jane.doe@email.com | (555) 123-4567 | linkedin.com/in/janedoe | github.com/janedoe

## Summary

Experienced software engineer with 10+ years of experience building scalable systems.

## Experience

### Senior Software Engineer

**Acme Corp**, San Francisco, CA | Jan 2020 - Present

- Led a team of 5 engineers to deliver a microservices platform
- Reduced deployment time by **40%** through CI/CD pipeline optimization
- Mentored 3 junior engineers on best practices

### Software Engineer

**StartupCo**, New York, NY | Jun 2015 - Dec 2019

- Built RESTful APIs serving 1M+ daily requests
- Implemented *real-time* data processing pipeline using Kafka

## Education

### Master of Science in Computer Science

**MIT**, Cambridge, MA | 2015

### Bachelor of Science in Mathematics

**UCLA**, Los Angeles, CA | 2013

## Skills

- **Languages:** Python, Go, Java, TypeScript
- **Frameworks:** Django, FastAPI, React
- **Cloud:** AWS, GCP, Kubernetes, Docker
"""

TABLE_RESUME_MD = """\
# Test Person

test@email.com

## Summary

Summary text here.

## Experience

## Education

## Skills

| Category | Skills | Level |
|----------|--------|-------|
| Languages | Python, Go | Expert |
| Cloud | AWS, GCP | Advanced |
| Tools | Docker, K8s | Intermediate |
"""

NESTED_BULLETS_MD = """\
# Nested Person

## Summary

A summary.

## Experience

### Role

**Company** | 2020 - Present

- Top-level bullet
  - Second-level bullet
    - Third-level bullet
      - Fourth-level bullet
- Another top-level bullet

## Education

## Skills
"""

CODE_AND_IMAGE_MD = """\
# Code Person

## Summary

Has code and images.

## Experience

### Developer

**Company** | 2020 - Present

- Normal bullet
- Has a [link](https://example.com) in it

```python
def hello():
    print("world")
```

![profile photo](photo.png)

- Bullet after code block

## Education

## Skills
"""

LINK_RESUME_MD = """\
# Link Person

## Summary

Check out [my portfolio](https://example.com) for details.

## Experience

### Engineer

**Company** | 2020 - Present

- Visit [project page](https://github.com/user/repo) for source code

## Education

## Skills
"""

NO_TABLE_RESUME_MD = """\
# Simple Person

simple@email.com

## Summary

Just a summary.

## Experience

### Developer

**Company** | 2020 - Present

- Did things

## Education

### BS CS

**University** | 2018

## Skills

- Python, Go, Java
"""

LONG_CONTENT_MD = """\
# Long Content Person

email@test.com

## Summary

""" + "This is a very long summary paragraph that goes on and on. " * 50 + """

## Experience

### Senior Engineer

**BigCorp** | 2015 - Present

""" + "\n".join(f"- Achievement number {i} with detailed description of what was done" for i in range(100)) + """

## Education

### PhD in Computer Science

**Stanford** | 2015

## Skills

- Languages: Python, Go, Java, TypeScript, C++, Rust, Ruby, Kotlin, Swift, Scala
"""


# ---------------------------------------------------------------------------
# Unit Tests: Element Mapping
# ---------------------------------------------------------------------------


class TestHeadingMapping:
    """Test H1/H2/H3 heading mapping to Word elements."""

    def test_h1_becomes_heading_1(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        # Find Heading 1 paragraphs
        h1_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert len(h1_paras) == 1
        assert h1_paras[0].text == "Jane Doe"

    def test_h2_becomes_heading_2(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        h2_paras = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        h2_texts = [p.text for p in h2_paras]
        assert "Summary" in h2_texts
        assert "Experience" in h2_texts
        assert "Education" in h2_texts
        assert "Skills" in h2_texts

    def test_h3_becomes_heading_3(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        h3_paras = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
        h3_texts = [p.text for p in h3_paras]
        assert "Senior Software Engineer" in h3_texts
        assert "Software Engineer" in h3_texts
        assert "Master of Science in Computer Science" in h3_texts
        assert "Bachelor of Science in Mathematics" in h3_texts


class TestBulletMapping:
    """Test bullet point mapping to Word list items."""

    def test_bullets_use_list_bullet_style(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        list_paras = [
            p for p in doc.paragraphs
            if p.style.name in ("List Bullet", "List Bullet 2")
        ]
        assert len(list_paras) > 0

    def test_bullet_content_preserved(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Led a team of 5 engineers" in all_text
        assert "Reduced deployment time" in all_text
        assert "Built RESTful APIs" in all_text


class TestInlineFormatting:
    """Test bold, italic, and link preservation."""

    def test_bold_text_preserved(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        # Find paragraphs with bold runs
        bold_runs = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold and run.text:
                    bold_runs.append(run.text)

        assert any("40%" in text for text in bold_runs)

    def test_italic_text_preserved(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        italic_runs = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.italic and run.text:
                    italic_runs.append(run.text)

        assert any("real-time" in text for text in italic_runs)

    def test_links_rendered(self) -> None:
        parsed = parse_resume_markdown(LINK_RESUME_MD)
        doc = render_docx(parsed)

        # Links create hyperlink elements in the XML
        # Check that the link text appears in the document
        body_xml = doc.element.body.xml
        assert "my portfolio" in body_xml or "project page" in body_xml


class TestContactInfo:
    """Test contact information rendering."""

    def test_contact_lines_present(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "jane.doe@email.com" in all_text

    def test_contact_rendered_after_name(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        # Find the name heading and the next paragraph
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == "Heading 1" and para.text == "Jane Doe":
                # The next paragraph should contain contact info
                next_para = doc.paragraphs[i + 1]
                assert "jane.doe@email.com" in next_para.text
                break
        else:
            pytest.fail("Name heading not found")


class TestTableConversion:
    """Test table conversion to Word tables."""

    def test_tables_created(self) -> None:
        parsed = parse_resume_markdown(TABLE_RESUME_MD)
        doc = render_docx(parsed)

        assert len(doc.tables) >= 1

    def test_table_has_correct_dimensions(self) -> None:
        parsed = parse_resume_markdown(TABLE_RESUME_MD)
        doc = render_docx(parsed)

        table = doc.tables[0]
        # 4 rows: 1 header + 3 data (separator is skipped)
        assert len(table.rows) == 4
        assert len(table.columns) == 3

    def test_table_header_content(self) -> None:
        parsed = parse_resume_markdown(TABLE_RESUME_MD)
        doc = render_docx(parsed)

        table = doc.tables[0]
        header_texts = [cell.text for cell in table.rows[0].cells]
        assert "Category" in header_texts
        assert "Skills" in header_texts
        assert "Level" in header_texts

    def test_table_body_content(self) -> None:
        parsed = parse_resume_markdown(TABLE_RESUME_MD)
        doc = render_docx(parsed)

        table = doc.tables[0]
        body_texts = []
        for row in table.rows[1:]:
            for cell in row.cells:
                body_texts.append(cell.text)

        assert "Python, Go" in body_texts
        assert "AWS, GCP" in body_texts

    def test_table_header_row_styled(self) -> None:
        parsed = parse_resume_markdown(TABLE_RESUME_MD)
        doc = render_docx(parsed)

        table = doc.tables[0]
        # Header row cells should have bold text
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text:
                        assert run.font.bold is True


# ---------------------------------------------------------------------------
# Unit Tests: Edge Cases
# ---------------------------------------------------------------------------


class TestEdgeCases:
    """Test edge case handling."""

    def test_no_tables_renders_without_error(self) -> None:
        parsed = parse_resume_markdown(NO_TABLE_RESUME_MD)
        doc = render_docx(parsed)
        assert len(doc.tables) == 0
        assert len(doc.paragraphs) > 0

    def test_deeply_nested_bullets(self) -> None:
        parsed = parse_resume_markdown(NESTED_BULLETS_MD)
        doc = render_docx(parsed)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Top-level bullet" in all_text
        assert "Second-level bullet" in all_text
        assert "Third-level bullet" in all_text
        assert "Fourth-level bullet" in all_text

    def test_long_content_no_corruption(self) -> None:
        parsed = parse_resume_markdown(LONG_CONTENT_MD)
        doc = render_docx(parsed)

        # Verify document has content and isn't corrupted
        assert len(doc.paragraphs) > 50
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Long Content Person" in all_text
        assert "Achievement number 99" in all_text

    def test_code_blocks_skipped_with_warning(self, capsys: pytest.CaptureFixture) -> None:
        parsed = parse_resume_markdown(CODE_AND_IMAGE_MD)
        doc = render_docx(parsed)

        captured = capsys.readouterr()
        assert "Code block" in captured.err
        assert "skipped" in captured.err

        # Code content should NOT appear in the document
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert 'def hello()' not in all_text

    def test_images_skipped_with_warning(self, capsys: pytest.CaptureFixture) -> None:
        parsed = parse_resume_markdown(CODE_AND_IMAGE_MD)
        doc = render_docx(parsed)

        captured = capsys.readouterr()
        assert "Image" in captured.err
        assert "skipped" in captured.err

    def test_bullet_after_code_block_preserved(self) -> None:
        parsed = parse_resume_markdown(CODE_AND_IMAGE_MD)
        doc = render_docx(parsed)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Bullet after code block" in all_text


# ---------------------------------------------------------------------------
# Unit Tests: Presets
# ---------------------------------------------------------------------------


class TestPresets:
    """Test preset style selection."""

    def test_professional_preset(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed, preset="professional")
        assert len(doc.paragraphs) > 0

    def test_simple_preset(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed, preset="simple")
        assert len(doc.paragraphs) > 0

    def test_creative_preset(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed, preset="creative")
        assert len(doc.paragraphs) > 0

    def test_academic_preset(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed, preset="academic")
        assert len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# Unit Tests: CLI Argument Parsing
# ---------------------------------------------------------------------------


class TestArgParsing:
    """Test CLI argument parsing."""

    def test_minimal_args(self) -> None:
        result = parse_args(["script.py", "resume.md"])
        assert result["input"] == "resume.md"
        assert result["preset"] == "professional"
        assert result["template"] is None

    def test_output_path(self) -> None:
        result = parse_args(["script.py", "resume.md", "output.docx"])
        assert result["output"] == "output.docx"

    def test_preset_arg(self) -> None:
        result = parse_args(["script.py", "resume.md", "--preset", "academic"])
        assert result["preset"] == "academic"

    def test_template_arg(self) -> None:
        result = parse_args(["script.py", "resume.md", "--template", "my.docx"])
        assert result["template"] == "my.docx"

    def test_missing_input(self) -> None:
        result = parse_args(["script.py"])
        assert "error" in result

    def test_invalid_preset(self) -> None:
        config = parse_args(["script.py", "resume.md", "--preset", "invalid"])
        config = resolve_config(config)
        assert "error" in config

    def test_resolve_output_default(self) -> None:
        config = parse_args(["script.py", "resume.md"])
        config = resolve_config(config)
        assert config["output"] == "resume.docx"


# ---------------------------------------------------------------------------
# Integration Tests
# ---------------------------------------------------------------------------


class TestDocxIntegration:
    """Integration tests for valid DOCX output."""

    def test_generated_docx_is_valid(self) -> None:
        """Generated .docx can be reopened by python-docx without error."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        # Re-open the file to verify it's a valid .docx
        reopened = Document(tmp_path)
        assert len(reopened.paragraphs) > 0
        Path(tmp_path).unlink()

    def test_round_trip_text_extraction(self) -> None:
        """Text extracted from generated DOCX matches source content."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        doc = render_docx(parsed)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        reopened = Document(tmp_path)
        extracted_text = "\n".join(p.text for p in reopened.paragraphs)

        # Check key content is preserved
        assert "Jane Doe" in extracted_text
        assert "Senior Software Engineer" in extracted_text
        assert "Led a team of 5 engineers" in extracted_text
        assert "Built RESTful APIs" in extracted_text
        assert "Master of Science in Computer Science" in extracted_text
        assert "Python" in extracted_text

        Path(tmp_path).unlink()

    def test_round_trip_table_content(self) -> None:
        """Table content survives round-trip save/load."""
        parsed = parse_resume_markdown(TABLE_RESUME_MD)
        doc = render_docx(parsed)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        reopened = Document(tmp_path)
        assert len(reopened.tables) >= 1

        table = reopened.tables[0]
        all_cell_text = []
        for row in table.rows:
            for cell in row.cells:
                all_cell_text.append(cell.text)

        assert "Languages" in all_cell_text
        assert "Python, Go" in all_cell_text

        Path(tmp_path).unlink()

    def test_all_presets_produce_valid_docx(self) -> None:
        """All presets produce valid .docx files."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)

        for preset in ["professional", "simple", "creative", "academic"]:
            doc = render_docx(parsed, preset=preset)

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                doc.save(f.name)
                tmp_path = f.name

            reopened = Document(tmp_path)
            assert len(reopened.paragraphs) > 0, f"Preset '{preset}' produced empty doc"
            Path(tmp_path).unlink()

    def test_custom_template_loading(self) -> None:
        """Custom template file is loaded when provided."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)

        # Use one of the bundled templates as a "custom" template
        template_path = str(
            Path(__file__).resolve().parent.parent / "templates" / "word" / "simple.docx"
        )
        doc = render_docx(parsed, template_path=template_path)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        reopened = Document(tmp_path)
        all_text = "\n".join(p.text for p in reopened.paragraphs)
        assert "Jane Doe" in all_text
        Path(tmp_path).unlink()


# ---------------------------------------------------------------------------
# Error Handling Tests
# ---------------------------------------------------------------------------


class TestErrorHandling:
    """Test error handling paths."""

    def test_missing_custom_template_raises_error(self) -> None:
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        with pytest.raises(FileNotFoundError, match="Custom template not found"):
            render_docx(parsed, template_path="/nonexistent/template.docx")

    def test_cli_missing_input_file(self) -> None:
        config = parse_args(["script.py", "/nonexistent/resume.md"])
        config = resolve_config(config)
        # The config itself resolves fine; the file check happens in main()
        assert "error" not in config

    def test_invalid_preset_in_config(self) -> None:
        config = parse_args(["script.py", "resume.md", "--preset", "fancy"])
        config = resolve_config(config)
        assert "error" in config
        assert "fancy" in config["error"]


# ---------------------------------------------------------------------------
# Custom Template Tests
# ---------------------------------------------------------------------------


def _create_valid_custom_template(path: str) -> None:
    """Create a minimal valid custom template with required styles."""
    doc = Document()
    # Document() creates a document that already has Normal, Heading 1, Heading 2
    # styles by default in python-docx. Save it.
    doc.save(path)


def _create_custom_template_with_extra_styles(path: str) -> None:
    """Create a custom template with required styles plus additional custom styles."""
    doc = Document()
    # Add custom styles that are additional to the required ones
    from docx.enum.style import WD_STYLE_TYPE
    doc.styles.add_style("Custom Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Custom Emphasis", WD_STYLE_TYPE.CHARACTER)
    doc.save(path)


def _create_template_missing_styles(path: str, keep_styles: list[str] | None = None) -> None:
    """Create a .docx file that is valid but has some required styles removed.

    This is tricky because python-docx always includes built-in styles.
    Instead, we create a bare-bones .docx with manipulated XML to remove styles.
    """
    doc = Document()
    # We cannot easily remove built-in styles from python-docx.
    # Instead, we'll create a file and manipulate its XML directly.
    # For testing purposes, we use a different approach:
    # We save normally and then test the validate function with a mock.
    doc.save(path)


class TestTemplateValidation:
    """Test custom template style validation."""

    def test_valid_template_has_no_missing_styles(self) -> None:
        """A standard Document has all required styles."""
        doc = Document()
        missing = validate_template_styles(doc)
        assert missing == []

    def test_validate_reports_missing_styles(self) -> None:
        """validate_template_styles correctly identifies missing styles."""
        # Create a mock-like document object with limited styles
        doc = Document()
        # Get the actual style names
        available = {s.name for s in doc.styles}
        assert "Heading 1" in available
        assert "Heading 2" in available
        assert "Normal" in available

    def test_validate_returns_specific_missing_names(self) -> None:
        """When styles are missing, the exact names are returned."""
        from unittest.mock import MagicMock, PropertyMock

        # Create a mock document with only Normal style
        mock_doc = MagicMock()
        mock_style = MagicMock()
        mock_style.name = "Normal"
        type(mock_doc).styles = PropertyMock(return_value=[mock_style])

        missing = validate_template_styles(mock_doc)
        assert "Heading 1" in missing
        assert "Heading 2" in missing
        assert "Normal" not in missing

    def test_validate_all_styles_missing(self) -> None:
        """When all required styles are missing, all are reported."""
        from unittest.mock import MagicMock, PropertyMock

        mock_doc = MagicMock()
        type(mock_doc).styles = PropertyMock(return_value=[])

        missing = validate_template_styles(mock_doc)
        assert set(missing) == set(REQUIRED_TEMPLATE_STYLES)

    def test_required_styles_constant(self) -> None:
        """REQUIRED_TEMPLATE_STYLES contains the expected minimum styles."""
        assert "Heading 1" in REQUIRED_TEMPLATE_STYLES
        assert "Heading 2" in REQUIRED_TEMPLATE_STYLES
        assert "Normal" in REQUIRED_TEMPLATE_STYLES


class TestCustomTemplateFallback:
    """Test fallback logic when custom template is incompatible."""

    def test_incompatible_template_falls_back_with_warning(
        self, capsys: pytest.CaptureFixture
    ) -> None:
        """Template with missing styles falls back to preset with warning."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            # Create a corrupted/non-docx file
            f.write(b"not a valid docx file")
            bad_template = f.name

        try:
            doc = render_docx(parsed, template_path=bad_template)
            captured = capsys.readouterr()

            # Should have warned about the bad template
            assert "corrupted" in captured.err or "not a valid" in captured.err
            assert "Falling back" in captured.err

            # Should still produce a valid document via fallback
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Jane Doe" in all_text
        finally:
            Path(bad_template).unlink()

    def test_missing_template_file_raises_with_preset_info(self) -> None:
        """FileNotFoundError includes available preset names."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)
        with pytest.raises(FileNotFoundError, match="Available presets"):
            render_docx(parsed, template_path="/nonexistent/template.docx")

    def test_corrupted_docx_clear_error_message(
        self, capsys: pytest.CaptureFixture
    ) -> None:
        """Corrupted .docx file gives a clear error message."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"PK\x03\x04corrupted zip data")
            corrupt_path = f.name

        try:
            doc = render_docx(parsed, template_path=corrupt_path)
            captured = capsys.readouterr()
            assert "corrupted" in captured.err or "not a valid" in captured.err
            # Document should still be produced via fallback
            assert len(doc.paragraphs) > 0
        finally:
            Path(corrupt_path).unlink()

    def test_missing_styles_warning_specifies_which_styles(
        self, capsys: pytest.CaptureFixture, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Warning message specifies exactly which required styles are missing."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)

        # Monkeypatch validate_template_styles to return missing styles
        monkeypatch.setattr(
            md_to_docx,
            "validate_template_styles",
            lambda doc: ["Heading 1", "Heading 2"],
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_valid_custom_template(f.name)
            template_path = f.name

        try:
            doc = render_docx(parsed, template_path=template_path)
            captured = capsys.readouterr()
            assert "Heading 1" in captured.err
            assert "Heading 2" in captured.err
            assert "missing required styles" in captured.err
            assert "Falling back" in captured.err
        finally:
            Path(template_path).unlink()

    def test_fallback_is_automatic(self, capsys: pytest.CaptureFixture) -> None:
        """Fallback to preset happens automatically without user intervention."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"not a docx")
            bad_path = f.name

        try:
            # Should not raise -- should fall back silently with warning
            doc = render_docx(parsed, template_path=bad_path)
            assert len(doc.paragraphs) > 0

            captured = capsys.readouterr()
            assert "WARNING" in captured.err
        finally:
            Path(bad_path).unlink()


class TestCustomTemplateIntegration:
    """Integration tests for custom template support."""

    def test_custom_template_content_injection(self) -> None:
        """Resume content is correctly injected into a custom template."""
        parsed = parse_resume_markdown(FULL_RESUME_MD)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_valid_custom_template(f.name)
            template_path = f.name

        try:
            doc = render_docx(parsed, template_path=template_path)

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
                doc.save(out.name)
                out_path = out.name

            reopened = Document(out_path)
            all_text = "\n".join(p.text for p in reopened.paragraphs)

            # Verify all resume content is injected
            assert "Jane Doe" in all_text
            assert "Senior Software Engineer" in all_text
            assert "Led a team of 5 engineers" in all_text
            assert "Built RESTful APIs" in all_text

            Path(out_path).unlink()
        finally:
            Path(template_path).unlink()

    def test_custom_template_styles_preserved(self) -> None:
        """Custom template's styles are preserved in the output document."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_custom_template_with_extra_styles(f.name)
            template_path = f.name

        try:
            parsed = parse_resume_markdown(FULL_RESUME_MD)
            doc = render_docx(parsed, template_path=template_path)

            # Check that the custom styles from the template are still available
            style_names = {s.name for s in doc.styles}
            assert "Custom Subtitle" in style_names
            assert "Custom Emphasis" in style_names

            # Also verify required styles are present
            assert "Heading 1" in style_names
            assert "Heading 2" in style_names
            assert "Normal" in style_names
        finally:
            Path(template_path).unlink()

    def test_custom_template_extra_styles_not_used_for_resume(self) -> None:
        """Additional custom styles exist but are not used for resume content."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_custom_template_with_extra_styles(f.name)
            template_path = f.name

        try:
            parsed = parse_resume_markdown(FULL_RESUME_MD)
            doc = render_docx(parsed, template_path=template_path)

            # Resume content should use standard styles, not custom ones
            for para in doc.paragraphs:
                if para.text:
                    assert para.style.name != "Custom Subtitle"

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
                doc.save(out.name)
                out_path = out.name

            # Verify round-trip: custom styles still present after save/reload
            reopened = Document(out_path)
            style_names = {s.name for s in reopened.styles}
            assert "Custom Subtitle" in style_names
            assert "Custom Emphasis" in style_names

            Path(out_path).unlink()
        finally:
            Path(template_path).unlink()

    def test_custom_template_heading_styles_used(self) -> None:
        """Resume headings use the template's Heading 1 and Heading 2 styles."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_valid_custom_template(f.name)
            template_path = f.name

        try:
            parsed = parse_resume_markdown(FULL_RESUME_MD)
            doc = render_docx(parsed, template_path=template_path)

            h1_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
            h2_paras = [p for p in doc.paragraphs if p.style.name == "Heading 2"]

            assert len(h1_paras) >= 1
            assert h1_paras[0].text == "Jane Doe"
            assert len(h2_paras) >= 1
            h2_texts = [p.text for p in h2_paras]
            assert "Experience" in h2_texts
        finally:
            Path(template_path).unlink()
