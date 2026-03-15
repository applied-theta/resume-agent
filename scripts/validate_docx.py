"""Post-export ATS validation for DOCX files.

Extracts text from a generated DOCX file using python-docx and verifies
that content round-trips correctly against the source markdown. Checks
that all section headings are present and no content was silently dropped.

Usage as module:
    from validate_docx import validate_docx

    result = validate_docx("output.docx", source_markdown)
    # result["passed"] -> bool
    # result["issues"] -> list[str]
    # result["missing_sections"] -> list[str]
    # result["summary"] -> str

Usage as CLI:
    uv run scripts/validate_docx.py <docx-path> <source-markdown-path>
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print(
        "Error: python-docx is not installed.\n"
        "Install it with: uv add python-docx\n"
        "Or: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)

from parse_resume import parse_resume_markdown


def extract_docx_text(docx_path: str) -> dict:
    """Extract all text content from a DOCX file.

    Returns a dict with:
        paragraphs: list[str] - all paragraph texts
        headings: list[str] - heading paragraph texts
        table_texts: list[str] - all cell texts from tables
        full_text: str - concatenation of all text content

    Args:
        docx_path: Path to the DOCX file.

    Raises:
        FileNotFoundError: If the DOCX file does not exist.
        ValueError: If the file cannot be opened as a valid DOCX.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(
            f"Cannot open DOCX file (corrupted or invalid format): {exc}"
        ) from exc

    paragraphs: list[str] = []
    headings: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            headings.append(text)

    table_texts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    table_texts.append(cell_text)

    all_parts = paragraphs + table_texts
    full_text = "\n".join(all_parts)

    return {
        "paragraphs": paragraphs,
        "headings": headings,
        "table_texts": table_texts,
        "full_text": full_text,
    }


def _normalize(text: str) -> str:
    """Normalize text for comparison: lowercase, collapse whitespace, strip markdown."""
    text = text.lower().strip()
    # Remove markdown bold/italic markers
    text = re.sub(r"\*+", "", text)
    # Collapse whitespace
    text = re.sub(r"\s+", " ", text)
    return text


def _content_present(needle: str, haystack: str) -> bool:
    """Check if content is present in the full text, using normalized comparison."""
    return _normalize(needle) in _normalize(haystack)


def validate_docx(docx_path: str, source_markdown: str) -> dict:
    """Validate a generated DOCX against its source markdown.

    Checks:
    1. All section headings from source markdown appear in the DOCX.
    2. No content was silently dropped (significant text lines are present).

    Args:
        docx_path: Path to the generated DOCX file.
        source_markdown: The original markdown content that was used to generate the DOCX.

    Returns:
        A dict with:
            passed: bool - overall validation result
            issues: list[str] - specific validation problems found
            missing_sections: list[str] - section headings not found in the DOCX
            dropped_content: list[str] - significant content lines not found
            summary: str - human-readable summary
    """
    issues: list[str] = []
    missing_sections: list[str] = []
    dropped_content: list[str] = []

    # Extract text from DOCX
    extracted = extract_docx_text(docx_path)
    full_text = extracted["full_text"]

    # Parse source markdown to get expected structure
    parsed = parse_resume_markdown(source_markdown)

    # Check 1: Verify candidate name is present
    if parsed["name"] and not _content_present(parsed["name"], full_text):
        issues.append(f"Candidate name missing: '{parsed['name']}'")

    # Check 2: Verify all section headings from source markdown are present
    for section in parsed["sections"]:
        heading = section["heading"]
        if not _content_present(heading, full_text):
            missing_sections.append(heading)
            issues.append(f"Section heading missing: '{heading}'")

    # Check 3: Check for silently dropped content
    # Check significant content lines (non-empty, non-heading, non-blank lines)
    for section in parsed["sections"]:
        heading = section["heading"]
        content_lines = section.get("content", [])

        for line in content_lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Skip markdown table separators (|---|---|)
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue

            # Skip code fence markers
            if stripped.startswith("```") or stripped.startswith("~~~"):
                continue

            # Skip image lines (images are intentionally skipped by the renderer)
            if re.match(r"!\[.*\]\(.*\)", stripped):
                continue

            # Extract meaningful text from the line for comparison
            # Strip bullet markers
            check_text = stripped
            if check_text.startswith("- ") or check_text.startswith("* "):
                check_text = check_text[2:].strip()

            # Strip markdown table delimiters for table row content
            if check_text.startswith("|") and check_text.endswith("|"):
                cells = [c.strip() for c in check_text.strip("|").split("|")]
                # Check each cell individually
                cells_found = True
                for cell in cells:
                    if cell and not _content_present(cell, full_text):
                        cells_found = False
                        break
                if not cells_found:
                    dropped_content.append(f"Table row in '{heading}': {stripped}")
                    issues.append(
                        f"Content dropped in '{heading}': table row '{stripped[:60]}...'"
                        if len(stripped) > 60
                        else f"Content dropped in '{heading}': table row '{stripped}'"
                    )
                continue

            # For regular content, skip very short fragments (< 4 chars)
            # as they may be formatting artifacts
            if len(check_text) < 4:
                continue

            # Strip inline markdown for comparison
            # Remove link syntax: [text](url) -> text
            plain_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", check_text)
            # Remove bold/italic markers
            plain_text = re.sub(r"\*+([^*]+)\*+", r"\1", plain_text)
            # Remove inline code
            plain_text = re.sub(r"`([^`]+)`", r"\1", plain_text)

            if not _content_present(plain_text, full_text):
                dropped_content.append(f"In '{heading}': {stripped}")
                issues.append(
                    f"Content dropped in '{heading}': '{stripped[:60]}...'"
                    if len(stripped) > 60
                    else f"Content dropped in '{heading}': '{stripped}'"
                )

    # Check meta lines
    for section in parsed["sections"]:
        meta_line = section.get("meta_line")
        if meta_line:
            plain_meta = re.sub(r"\*+([^*]+)\*+", r"\1", meta_line)
            if not _content_present(plain_meta, full_text):
                dropped_content.append(f"Meta line in '{section['heading']}': {meta_line}")
                issues.append(
                    f"Meta line dropped in '{section['heading']}': '{meta_line[:60]}'"
                )

    # Check contact lines
    for contact in parsed["contact_lines"]:
        # Extract email/url parts for checking (contact lines may be reformatted)
        plain_contact = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", contact)
        if not _content_present(plain_contact, full_text):
            dropped_content.append(f"Contact: {contact}")
            issues.append(f"Contact information dropped: '{contact[:60]}'")

    passed = len(issues) == 0

    # Build summary with enhanced missing sections report
    if passed:
        section_count = len(parsed["sections"])
        summary = (
            f"DOCX validation passed. "
            f"All {section_count} sections and content verified."
        )
    else:
        from export_edge_cases import format_missing_sections_report
        total_sections = len(parsed["sections"])
        if missing_sections:
            sections_report = format_missing_sections_report(
                missing_sections, total_sections
            )
        else:
            sections_report = ""

        summary = (
            f"DOCX validation failed with {len(issues)} issue(s): "
            f"{len(missing_sections)} missing section(s), "
            f"{len(dropped_content)} dropped content item(s)."
        )
        if sections_report:
            summary += f"\n\n{sections_report}"

    return {
        "passed": passed,
        "issues": issues,
        "missing_sections": missing_sections,
        "dropped_content": dropped_content,
        "summary": summary,
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> int:
    """CLI entry point for DOCX ATS validation."""
    if len(sys.argv) < 3:
        print(
            f"Usage: {sys.argv[0]} <docx-path> <source-markdown-path>",
            file=sys.stderr,
        )
        return 1

    docx_path = sys.argv[1]
    md_path = sys.argv[2]

    # Read source markdown
    md_file = Path(md_path)
    if not md_file.exists():
        print(f"Error: Source markdown not found: {md_path}", file=sys.stderr)
        return 1

    try:
        source_md = md_file.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        print(f"Error: Cannot read markdown file: {exc}", file=sys.stderr)
        return 1

    # Validate
    try:
        result = validate_docx(docx_path, source_md)
    except FileNotFoundError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    # Output result
    if result["passed"]:
        print(f"PASS: {result['summary']}")
        return 0
    else:
        print(f"FAIL: {result['summary']}")
        for issue in result["issues"]:
            print(f"  - {issue}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
