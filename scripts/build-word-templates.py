"""Build bundled Word preset template files for DOCX export.

Generates 4 .docx template files in templates/word/ with pre-configured
heading styles, body text, list styles, and table styles appropriate for
each use case.

Usage:
    uv run scripts/build-word-templates.py

Presets:
    professional  Calibri, Navy (#2B547E) - Corporate and business roles
    simple        Arial, Black (#000000) - Maximum ATS compatibility
    creative      Georgia, Teal (#008080) - Design, marketing, creative roles
    academic      Times New Roman, Dark Gray (#333333) - Research, education, CV-style
"""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
OUTPUT_DIR = PROJECT_ROOT / "templates" / "word"

PRESETS: dict[str, dict] = {
    "professional": {
        "font": "Calibri",
        "heading_color": RGBColor(0x2B, 0x54, 0x7E),  # Navy
        "body_font_size": Pt(11),
        "h1_font_size": Pt(22),
        "h2_font_size": Pt(13),
        "h3_font_size": Pt(11),
        "line_spacing": Pt(14),
        "space_after_body": Pt(4),
        "space_before_h2": Pt(12),
        "space_after_h2": Pt(4),
        "space_before_h3": Pt(8),
        "space_after_h3": Pt(2),
        "margin_top": Inches(0.75),
        "margin_bottom": Inches(0.75),
        "margin_left": Inches(0.75),
        "margin_right": Inches(0.75),
        "table_header_bg": RGBColor(0x2B, 0x54, 0x7E),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x2B, 0x54, 0x7E),
        "description": "Corporate and business roles",
    },
    "simple": {
        "font": "Arial",
        "heading_color": RGBColor(0x00, 0x00, 0x00),  # Black
        "body_font_size": Pt(11),
        "h1_font_size": Pt(20),
        "h2_font_size": Pt(13),
        "h3_font_size": Pt(11),
        "line_spacing": Pt(14),
        "space_after_body": Pt(4),
        "space_before_h2": Pt(12),
        "space_after_h2": Pt(4),
        "space_before_h3": Pt(8),
        "space_after_h3": Pt(2),
        "margin_top": Inches(1.0),
        "margin_bottom": Inches(1.0),
        "margin_left": Inches(1.0),
        "margin_right": Inches(1.0),
        "table_header_bg": RGBColor(0x00, 0x00, 0x00),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x00, 0x00, 0x00),
        "description": "Maximum ATS compatibility",
    },
    "creative": {
        "font": "Georgia",
        "heading_color": RGBColor(0x00, 0x80, 0x80),  # Teal
        "body_font_size": Pt(11),
        "h1_font_size": Pt(24),
        "h2_font_size": Pt(14),
        "h3_font_size": Pt(11),
        "line_spacing": Pt(15),
        "space_after_body": Pt(4),
        "space_before_h2": Pt(14),
        "space_after_h2": Pt(4),
        "space_before_h3": Pt(8),
        "space_after_h3": Pt(2),
        "margin_top": Inches(0.75),
        "margin_bottom": Inches(0.75),
        "margin_left": Inches(0.85),
        "margin_right": Inches(0.85),
        "table_header_bg": RGBColor(0x00, 0x80, 0x80),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x00, 0x80, 0x80),
        "description": "Design, marketing, creative roles",
    },
    "academic": {
        "font": "Times New Roman",
        "heading_color": RGBColor(0x33, 0x33, 0x33),  # Dark Gray
        "body_font_size": Pt(12),
        "h1_font_size": Pt(22),
        "h2_font_size": Pt(14),
        "h3_font_size": Pt(12),
        "line_spacing": Pt(16),
        "space_after_body": Pt(6),
        "space_before_h2": Pt(14),
        "space_after_h2": Pt(6),
        "space_before_h3": Pt(10),
        "space_after_h3": Pt(4),
        "margin_top": Inches(1.0),
        "margin_bottom": Inches(1.0),
        "margin_left": Inches(1.0),
        "margin_right": Inches(1.0),
        "table_header_bg": RGBColor(0x33, 0x33, 0x33),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x33, 0x33, 0x33),
        "description": "Research, education, CV-style",
    },
}


# ---------------------------------------------------------------------------
# Style configuration helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell: object, color: RGBColor) -> None:
    """Set the background shading color on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), f"{color}")
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(
    table: object,
    color: RGBColor,
    size: int = 4,
) -> None:
    """Set borders on a Word table using OxmlElement manipulation."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), f"{color}")
        borders.append(element)

    # Remove existing borders if present
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    tbl_pr.append(borders)


def configure_list_style(
    doc: Document,
    style_name: str,
    preset: dict,
    *,
    indent_left: Emu = Inches(0.25),
) -> None:
    """Configure or create a list bullet style in the document."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    font = style.font
    font.name = preset["font"]
    font.size = preset["body_font_size"]
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    pf = style.paragraph_format
    pf.left_indent = indent_left
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_heading_underline(paragraph: object, color: RGBColor) -> None:
    """Add a bottom border (underline) to a heading paragraph via XML."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color}")
    p_borders.append(bottom)
    p_pr.append(p_borders)


# ---------------------------------------------------------------------------
# Template builder
# ---------------------------------------------------------------------------


def build_template(name: str, preset: dict) -> Document:
    """Build a Word document template with the specified preset styles."""
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = preset["margin_top"]
        section.bottom_margin = preset["margin_bottom"]
        section.left_margin = preset["margin_left"]
        section.right_margin = preset["margin_right"]

    # -- Normal (body text) style --
    normal = doc.styles["Normal"]
    normal.font.name = preset["font"]
    normal.font.size = preset["body_font_size"]
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.space_after = preset["space_after_body"]
    normal.paragraph_format.line_spacing = preset["line_spacing"]
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # Set the rFonts element for font fallback (East Asian, complex script)
    rpr = normal._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), preset["font"])
    r_fonts.set(qn("w:hAnsi"), preset["font"])
    r_fonts.set(qn("w:cs"), preset["font"])

    # -- Heading 1 style (candidate name) --
    h1 = doc.styles["Heading 1"]
    h1.font.name = preset["font"]
    h1.font.size = preset["h1_font_size"]
    h1.font.bold = True
    h1.font.color.rgb = preset["heading_color"]
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # -- Heading 2 style (section headings) --
    h2 = doc.styles["Heading 2"]
    h2.font.name = preset["font"]
    h2.font.size = preset["h2_font_size"]
    h2.font.bold = True
    h2.font.color.rgb = preset["heading_color"]
    h2.paragraph_format.space_before = preset["space_before_h2"]
    h2.paragraph_format.space_after = preset["space_after_h2"]
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # -- Heading 3 style (subsection headings) --
    h3 = doc.styles["Heading 3"]
    h3.font.name = preset["font"]
    h3.font.size = preset["h3_font_size"]
    h3.font.bold = True
    h3.font.color.rgb = preset["heading_color"]
    h3.paragraph_format.space_before = preset["space_before_h3"]
    h3.paragraph_format.space_after = preset["space_after_h3"]
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # -- List Bullet style --
    configure_list_style(doc, "List Bullet", preset)

    # -- List Bullet 2 (nested bullets) --
    configure_list_style(
        doc,
        "List Bullet 2",
        preset,
        indent_left=Inches(0.5),
    )

    # -- Table style --
    # python-docx does not support creating new table styles directly via
    # the high-level API, so we configure the default "Table Grid" style.
    try:
        table_style = doc.styles["Table Grid"]
        table_style.font.name = preset["font"]
        table_style.font.size = preset["body_font_size"]
    except KeyError:
        pass

    # Create a sample table to embed table formatting in the template.
    # This ensures the table style definitions are stored in the .docx.
    # The renderer will use these styles when creating tables.
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = doc.styles["Table Grid"]

    # Apply borders
    set_table_borders(table, preset["table_border_color"])

    # Style header row
    for cell in table.rows[0].cells:
        set_cell_shading(cell, preset["table_header_bg"])
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = preset["table_header_text"]
                run.font.bold = True

    # Set header row cells with placeholder text to define formatting
    headers = ["Category", "Skills", "Level"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = preset["font"]
                run.font.size = preset["body_font_size"]
                run.font.color.rgb = preset["table_header_text"]
                run.font.bold = True

    # Style body row
    body_labels = ["", "", ""]
    for i, label in enumerate(body_labels):
        cell = table.rows[1].cells[i]
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = preset["font"]
                run.font.size = preset["body_font_size"]

    # Now remove the sample table content but keep the style definitions.
    # We clear the document body of paragraphs and tables, leaving only
    # the style definitions in the document.
    body = doc.element.body
    # Remove all child elements from the body (paragraphs, tables, etc.)
    for child in list(body):
        # Keep sectPr (section properties for margins)
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> int:
    """Generate all Word preset template files."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    for name, preset in PRESETS.items():
        doc = build_template(name, preset)
        output_path = OUTPUT_DIR / f"{name}.docx"
        doc.save(str(output_path))

        file_size = output_path.stat().st_size
        size_str = (
            f"{file_size / 1024:.1f} KB"
            if file_size < 1024 * 1024
            else f"{file_size / (1024 * 1024):.1f} MB"
        )
        print(f"Generated: {output_path.name} ({size_str}) - {preset['description']}")

    print(f"\nAll {len(PRESETS)} templates written to {OUTPUT_DIR}/")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
