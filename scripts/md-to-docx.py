"""Convert resume markdown to DOCX using python-docx.

Consumes the shared parser's structured resume data and produces .docx files
with proper heading hierarchy, bullet lists, inline formatting, tables,
and contact information.

Usage:
    uv run scripts/md-to-docx.py <input-md> [output-docx] [options]

Options:
    --preset PRESET      Style preset: professional, simple, creative, academic
                         (default: professional)
    --template PATH      Custom .docx template file (overrides preset)

Exits with code 0 on success, 1 on error. Errors are written to stderr.
"""

import re
import sys
import warnings as _warnings
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print(
        "Error: python-docx is not installed.\n"
        "Install it with: uv add python-docx\n"
        "Or: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)

from parse_resume import parse_resume_markdown

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
TEMPLATES_DIR = PROJECT_ROOT / "templates" / "word"

VALID_PRESETS = ["professional", "simple", "creative", "academic"]

# Required styles that custom templates must define.
REQUIRED_TEMPLATE_STYLES = ["Heading 1", "Heading 2", "Normal"]

# Fallback preset styling when no template file is available.
# These mirror the values used in build-word-templates.py.
PRESET_STYLES: dict[str, dict] = {
    "professional": {
        "font": "Calibri",
        "heading_color": RGBColor(0x2B, 0x54, 0x7E),
        "body_font_size": Pt(11),
        "h1_font_size": Pt(22),
        "h2_font_size": Pt(13),
        "h3_font_size": Pt(11),
        "table_header_bg": RGBColor(0x2B, 0x54, 0x7E),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x2B, 0x54, 0x7E),
    },
    "simple": {
        "font": "Arial",
        "heading_color": RGBColor(0x00, 0x00, 0x00),
        "body_font_size": Pt(11),
        "h1_font_size": Pt(20),
        "h2_font_size": Pt(13),
        "h3_font_size": Pt(11),
        "table_header_bg": RGBColor(0x00, 0x00, 0x00),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x00, 0x00, 0x00),
    },
    "creative": {
        "font": "Georgia",
        "heading_color": RGBColor(0x00, 0x80, 0x80),
        "body_font_size": Pt(11),
        "h1_font_size": Pt(24),
        "h2_font_size": Pt(14),
        "h3_font_size": Pt(11),
        "table_header_bg": RGBColor(0x00, 0x80, 0x80),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x00, 0x80, 0x80),
    },
    "academic": {
        "font": "Times New Roman",
        "heading_color": RGBColor(0x33, 0x33, 0x33),
        "body_font_size": Pt(12),
        "h1_font_size": Pt(22),
        "h2_font_size": Pt(14),
        "h3_font_size": Pt(12),
        "table_header_bg": RGBColor(0x33, 0x33, 0x33),
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_border_color": RGBColor(0x33, 0x33, 0x33),
    },
}


# ---------------------------------------------------------------------------
# Inline formatting helpers
# ---------------------------------------------------------------------------

# Regex patterns for inline markdown formatting
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_CODE_INLINE_RE = re.compile(r"`([^`]+)`")


def _add_inline_runs(paragraph: object, text: str, preset_styles: dict) -> None:
    """Parse markdown inline formatting and add runs to a paragraph.

    Handles **bold**, *italic*, [text](url) links, and `inline code`.
    Formatting can be nested (e.g., **bold *and italic* text**).
    """
    # Tokenize the text into segments: plain, bold, italic, link, code
    segments = _tokenize_inline(text)

    for seg_type, seg_text, seg_url in segments:
        if seg_type == "link":
            _add_hyperlink(paragraph, seg_url, seg_text, preset_styles)
        elif seg_type == "bold":
            run = paragraph.add_run(seg_text)
            run.bold = True
            run.font.name = preset_styles.get("font", "Calibri")
            run.font.size = preset_styles.get("body_font_size", Pt(11))
        elif seg_type == "italic":
            run = paragraph.add_run(seg_text)
            run.italic = True
            run.font.name = preset_styles.get("font", "Calibri")
            run.font.size = preset_styles.get("body_font_size", Pt(11))
        elif seg_type == "code":
            run = paragraph.add_run(seg_text)
            run.font.name = "Courier New"
            run.font.size = preset_styles.get("body_font_size", Pt(11))
        else:
            run = paragraph.add_run(seg_text)
            run.font.name = preset_styles.get("font", "Calibri")
            run.font.size = preset_styles.get("body_font_size", Pt(11))


def _tokenize_inline(text: str) -> list[tuple[str, str, str]]:
    """Tokenize text into inline formatting segments.

    Returns a list of (type, text, url) tuples where type is one of:
    'plain', 'bold', 'italic', 'link', 'code'.
    url is only populated for 'link' type.
    """
    segments: list[tuple[str, str, str]] = []
    pos = 0

    # Combined pattern to match any inline format token
    combined = re.compile(
        r"(?P<link>\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))"
        r"|(?P<bold>\*\*(?P<bold_text>.+?)\*\*)"
        r"|(?P<italic>(?<!\*)\*(?!\*)(?P<italic_text>.+?)(?<!\*)\*(?!\*))"
        r"|(?P<code>`(?P<code_text>[^`]+)`)"
    )

    for match in combined.finditer(text):
        start = match.start()
        if start > pos:
            segments.append(("plain", text[pos:start], ""))

        if match.group("link"):
            segments.append(("link", match.group("link_text"), match.group("link_url")))
        elif match.group("bold"):
            segments.append(("bold", match.group("bold_text"), ""))
        elif match.group("italic"):
            segments.append(("italic", match.group("italic_text"), ""))
        elif match.group("code"):
            segments.append(("code", match.group("code_text"), ""))

        pos = match.end()

    if pos < len(text):
        segments.append(("plain", text[pos:], ""))

    return segments


def _add_hyperlink(paragraph: object, url: str, text: str, preset_styles: dict) -> None:
    """Add a hyperlink to a paragraph using OxmlElement manipulation.

    python-docx does not have a high-level API for hyperlinks, so we
    construct the XML directly.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # Style as hyperlink (blue, underline)
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "0563C1")
    r_pr.append(color_elem)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    # Set font
    r_fonts = OxmlElement("w:rFonts")
    font_name = preset_styles.get("font", "Calibri")
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_pr.append(r_fonts)

    # Set font size
    font_size = preset_styles.get("body_font_size", Pt(11))
    sz = OxmlElement("w:sz")
    # python-docx Pt() returns EMU; Word XML needs half-points
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
    r_pr.append(sz)

    new_run.append(r_pr)
    new_run.text = text

    # Set the text node properly
    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = text
    new_run.append(t_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------


def _is_table_line(line: str) -> bool:
    """Check if a line is a markdown table row."""
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _is_table_separator(line: str) -> bool:
    """Check if a line is a markdown table separator (e.g., |---|---|)."""
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def _parse_table_cells(line: str) -> list[str]:
    """Extract cell contents from a markdown table row."""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _set_cell_shading(cell: object, color: RGBColor) -> None:
    """Set the background shading color on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), f"{color}")
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table: object, color: RGBColor, size: int = 4) -> None:
    """Set borders on a Word table using OxmlElement manipulation."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), f"{color}")
        borders.append(element)

    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    tbl_pr.append(borders)


def _add_table(doc: Document, table_lines: list[str], preset_styles: dict) -> None:
    """Convert markdown table lines to a Word table in the document."""
    rows_data: list[list[str]] = []
    for line in table_lines:
        if _is_table_separator(line):
            continue
        cells = _parse_table_cells(line)
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(row) for row in rows_data)
    num_rows = len(rows_data)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply table grid style if available
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass

    # Set borders
    _set_table_borders(table, preset_styles.get("table_border_color", RGBColor(0, 0, 0)))

    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.rows[row_idx].cells[col_idx]
            # Clear existing paragraphs
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_inline_runs(paragraph, cell_text, preset_styles)

            # Style header row (first row)
            if row_idx == 0:
                _set_cell_shading(
                    cell,
                    preset_styles.get("table_header_bg", RGBColor(0, 0, 0)),
                )
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = preset_styles.get(
                        "table_header_text", RGBColor(0xFF, 0xFF, 0xFF)
                    )


# ---------------------------------------------------------------------------
# Content rendering
# ---------------------------------------------------------------------------


def _get_bullet_level(line: str) -> int:
    """Determine the nesting level of a bullet line.

    Level 0: no indent or first-level bullet
    Level 1: 2-4 spaces indent
    Level 2+: 4+ spaces indent per level
    """
    stripped = line.lstrip()
    if not (stripped.startswith("- ") or stripped.startswith("* ")):
        return 0
    indent = len(line) - len(stripped)
    if indent < 2:
        return 0
    return min(indent // 2, 8)  # Cap at 8 levels for safety


def _get_bullet_text(line: str) -> str:
    """Extract the text content from a bullet line."""
    stripped = line.strip()
    if stripped.startswith("- "):
        return stripped[2:]
    if stripped.startswith("* "):
        return stripped[2:]
    return stripped


def _is_code_fence(line: str) -> bool:
    """Check if a line is a code fence (``` or ~~~)."""
    stripped = line.strip()
    return stripped.startswith("```") or stripped.startswith("~~~")


def _is_image_line(line: str) -> bool:
    """Check if a line is a markdown image (![alt](url))."""
    return bool(re.match(r"!\[.*\]\(.*\)", line.strip()))


def render_docx(parsed: dict, preset: str = "professional",
                template_path: str | None = None) -> Document:
    """Render parsed resume data into a Word document.

    Args:
        parsed: Output from parse_resume_markdown().
        preset: Style preset name (professional, simple, creative, academic).
        template_path: Optional path to a custom .docx template file.

    Returns:
        A python-docx Document object ready to be saved.
    """
    render_warnings: list[str] = []
    styles = PRESET_STYLES.get(preset, PRESET_STYLES["professional"])

    # Load document from template or create new
    doc = _load_template(preset, template_path, render_warnings)

    # -- Candidate name (H1 -> Heading 1) --
    if parsed["name"]:
        h1 = doc.add_heading(parsed["name"], level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- Contact information --
    if parsed["contact_lines"]:
        for contact_line in parsed["contact_lines"]:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_runs(para, contact_line, styles)

    # -- Sections --
    in_code_block = False

    for section in parsed["sections"]:
        heading = section["heading"]
        level = section["level"]

        # H2 -> Heading 2, H3 -> Heading 3
        doc.add_heading(heading, level=level)

        # Meta line (company/date for job entries)
        if section.get("meta_line"):
            meta_text = section["meta_line"]
            meta_para = doc.add_paragraph()
            _add_inline_runs(meta_para, meta_text, styles)

        # Content lines
        content_lines = section.get("content", [])
        idx = 0
        while idx < len(content_lines):
            line = content_lines[idx]
            stripped = line.strip()

            # Skip blank lines
            if not stripped:
                idx += 1
                continue

            # Handle code fences: skip with warning
            if _is_code_fence(stripped):
                if not in_code_block:
                    in_code_block = True
                    render_warnings.append(
                        f"Code block in section '{heading}' skipped (unsupported in DOCX export)"
                    )
                else:
                    in_code_block = False
                idx += 1
                continue

            if in_code_block:
                idx += 1
                continue

            # Handle images: skip with warning
            if _is_image_line(stripped):
                render_warnings.append(
                    f"Image in section '{heading}' skipped (unsupported in DOCX export)"
                )
                idx += 1
                continue

            # Handle markdown tables
            if _is_table_line(stripped):
                table_block: list[str] = []
                while idx < len(content_lines) and _is_table_line(
                    content_lines[idx].strip()
                ):
                    table_block.append(content_lines[idx])
                    idx += 1
                _add_table(doc, table_block, styles)
                continue

            # Handle bullet points
            if stripped.startswith("- ") or stripped.startswith("* "):
                bullet_text = _get_bullet_text(stripped)
                nest_level = _get_bullet_level(line)

                # Use appropriate list style based on nesting
                if nest_level == 0:
                    style_name = "List Bullet"
                elif nest_level == 1:
                    style_name = "List Bullet 2"
                else:
                    # For deeper nesting (3+ levels), use List Bullet 2
                    # with additional indent
                    style_name = "List Bullet 2"

                try:
                    para = doc.add_paragraph(style=style_name)
                except KeyError:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.25 * (nest_level + 1))

                _add_inline_runs(para, bullet_text, styles)

                # Handle deeply nested bullets by increasing indent
                if nest_level >= 2:
                    para.paragraph_format.left_indent = Inches(0.25 * (nest_level + 1))

                idx += 1
                continue

            # Regular paragraph text
            para = doc.add_paragraph()
            _add_inline_runs(para, stripped, styles)
            idx += 1

    # Print render warnings to stderr
    for warning in render_warnings:
        print(f"WARNING: {warning}", file=sys.stderr)

    return doc


def validate_template_styles(doc: Document) -> list[str]:
    """Validate that a template document contains the required styles.

    Args:
        doc: A python-docx Document loaded from a template file.

    Returns:
        A list of missing required style names. Empty list means valid.
    """
    available_styles = {style.name for style in doc.styles}
    missing = [s for s in REQUIRED_TEMPLATE_STYLES if s not in available_styles]
    return missing


def _load_template(preset: str, template_path: str | None,
                   render_warnings: list[str]) -> Document:
    """Load a document template from file or create a new document.

    Tries in order:
    1. Custom template path (if provided and valid)
    2. Bundled preset template
    3. New empty document

    Custom templates are validated for required styles (Heading 1, Heading 2,
    Normal). If validation fails, falls back to preset with a warning that
    specifies which styles are missing.
    """
    if template_path:
        custom_path = Path(template_path)
        if not custom_path.exists():
            raise FileNotFoundError(
                f"Custom template not found: {template_path}. "
                f"Available presets: {', '.join(VALID_PRESETS)}"
            )
        try:
            doc = Document(str(custom_path))
        except Exception as exc:
            render_warnings.append(
                f"Custom template '{template_path}' is corrupted or not a valid "
                f".docx file: {exc}. Falling back to preset '{preset}'."
            )
            doc = None

        if doc is not None:
            missing = validate_template_styles(doc)
            if missing:
                render_warnings.append(
                    f"Custom template '{template_path}' is missing required styles: "
                    f"{', '.join(missing)}. Required styles: "
                    f"{', '.join(REQUIRED_TEMPLATE_STYLES)}. "
                    f"Falling back to preset '{preset}'."
                )
            else:
                return doc

    # Try bundled preset template
    preset_template = TEMPLATES_DIR / f"{preset}.docx"
    if preset_template.exists():
        try:
            return Document(str(preset_template))
        except Exception:
            render_warnings.append(
                f"Bundled template for preset '{preset}' could not be loaded. "
                f"Using default document."
            )

    return Document()


# ---------------------------------------------------------------------------
# CLI argument parsing
# ---------------------------------------------------------------------------


def parse_args(argv: list[str]) -> dict:
    """Parse CLI arguments into a config dict."""
    if len(argv) < 2:
        return {"error": "Missing input file argument"}

    config: dict = {
        "input": argv[1],
        "output": None,
        "preset": "professional",
        "template": None,
    }

    i = 2
    while i < len(argv):
        arg = argv[i]
        if arg == "--preset" and i + 1 < len(argv):
            config["preset"] = argv[i + 1].lower()
            i += 2
        elif arg == "--template" and i + 1 < len(argv):
            config["template"] = argv[i + 1]
            i += 2
        elif not arg.startswith("--") and config["output"] is None:
            config["output"] = arg
            i += 1
        else:
            return {"error": f"Unknown argument: {arg}"}

    return config


def resolve_config(config: dict) -> dict:
    """Resolve config with defaults and validate values."""
    preset = config["preset"]
    if preset not in VALID_PRESETS:
        return {
            "error": f"Invalid preset: {preset}. "
            f"Valid presets: {', '.join(VALID_PRESETS)}"
        }

    # Resolve output path
    if config["output"] is None:
        input_path = Path(config["input"])
        config["output"] = str(input_path.with_suffix(".docx"))

    return config


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> int:
    """CLI entry point for DOCX generation."""
    config = parse_args(sys.argv)
    if "error" in config:
        print(f"Error: {config['error']}", file=sys.stderr)
        print(
            f"Usage: {sys.argv[0]} <input-md> [output-docx] "
            "[--preset PRESET] [--template PATH]",
            file=sys.stderr,
        )
        return 1

    config = resolve_config(config)
    if "error" in config:
        print(f"Error: {config['error']}", file=sys.stderr)
        return 1

    # Validate source file with actionable error message
    input_path = Path(config["input"])
    from export_edge_cases import check_source_file
    source_error = check_source_file(input_path)
    if source_error:
        print(f"Error: {source_error}", file=sys.stderr)
        return 1

    try:
        md_text = input_path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        print(f"Error: File is not valid text/markdown: {exc}", file=sys.stderr)
        return 1

    if not md_text.strip():
        print("Error: Input file is empty", file=sys.stderr)
        return 1

    # Parse markdown
    try:
        parsed = parse_resume_markdown(md_text)
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    # Print parser warnings
    for warning in parsed.get("warnings", []):
        print(f"WARNING: {warning}", file=sys.stderr)

    # Generate DOCX
    try:
        doc = render_docx(
            parsed,
            preset=config["preset"],
            template_path=config.get("template"),
        )
    except FileNotFoundError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    # Save document
    output_path = config["output"]
    doc.save(output_path)

    # Check output file size against Cowork limit
    from export_edge_cases import check_output_size, emit_warnings
    size_warnings = check_output_size(output_path)
    if size_warnings:
        emit_warnings(size_warnings)

    file_size = Path(output_path).stat().st_size
    size_str = (
        f"{file_size / 1024:.1f} KB"
        if file_size < 1024 * 1024
        else f"{file_size / (1024 * 1024):.1f} MB"
    )
    print(f"DOCX generated: {output_path} ({size_str})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
